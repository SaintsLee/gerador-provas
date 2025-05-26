import pandas as pd
import streamlit as st

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from io import BytesIO
from datetime import date

@st.cache_data
def ler_dataframe(dataframe, nome_aba):
    df = pd.read_excel(dataframe, sheet_name=nome_aba)
    return df

def monta_prova(dataframe, dict_prova):
    list_dfs = []

    for nome_aba, categorias_dict in dict_prova.items():
        df = ler_dataframe(dataframe, nome_aba)

        for categoria, qtd in categorias_dict.items():
            df_categoria = df[df['Tópico'] == categoria]
            total_disponivel = df_categoria.shape[0]

            if qtd > total_disponivel:
                raise ValueError(
                    f"{nome_aba} - Categoria '{categoria}': solicitado {qtd}, mas só há {total_disponivel} disponíveis.")

            questoes_sorteadas = df_categoria.sample(n=qtd)
            questoes_sorteadas = questoes_sorteadas.copy()
            questoes_sorteadas['Nivel'] = nome_aba  # opcional: marca o nível de origem
            list_dfs.append(questoes_sorteadas)

    prova = pd.concat(list_dfs).reset_index(drop=True)
    prova_embaralhada = prova.sample(frac=1).reset_index(drop=True)
    return prova_embaralhada

def listar_opcoes(abas, dataframe):
    lista_opcoes = []
    for aba in range(len(abas)):
        opcoes = pd.read_excel(dataframe,sheet_name=abas[aba])['Tópico'].value_counts().reset_index(name='Quantidade')
        lista_opcoes.append(opcoes)
    return lista_opcoes

def gera_opcoes(df,aba,com_reset=True):
    selecionados = {}

    # Botão de reset individual (dentro da aba)
    if com_reset and st.button(f"🔄 Resetar {aba}", key=f"reset_{aba}"):
        for _, row in df.iterrows():
            topico = row['Tópico']
            st.session_state[f"cb_{topico}_{aba}"] = False
            st.session_state[f"num_{topico}_{aba}"] = row['Quantidade']  # ou 0 se quiser zerar

    # Dicionário para armazenar os resultados
    for _, row in df.iterrows():
        topico = row['Tópico']
        quantidade_padrao = row['Quantidade']

        # Checkbox para selecionar o tópico
        if st.checkbox(f"{topico}", key=f"cb_{topico}_{aba}"):
            # Input para definir a quantidade, visível apenas se o tópico for selecionado
            quantidade_escolhida = st.number_input(
                f"Quantidade para {topico}:",
                min_value=0,
                max_value=quantidade_padrao,
                value=int(quantidade_padrao),
                step=1,
                key=f"num_{topico}_{aba}"
            )
            selecionados[topico] = quantidade_escolhida
    return selecionados

def gerar_dict_prova(abas,lista_questoes):
    prova_dict = {}
    for aba in range(len(abas)):
        prova_dict[abas[aba]] = lista_questoes[aba]
    return prova_dict

def montar_prova_doc(df, titulo):
    data = date.today().strftime("%d/%m/%y")

    doc = Document()
    doc.add_heading(titulo, 0)

    # Subtítulo com estilo 'Subtitle'
    sub = doc.add_paragraph(data)
    sub.style = 'Subtitle'

    for i, row in df.iterrows():
        # Cria o parágrafo
        par = doc.add_paragraph()

        # Adiciona o número em negrito
        run_num = par.add_run(f"{i + 1}. ")
        run_num.bold = True
        run_num.font.size = Pt(12)

        # Adiciona o texto da questão normal
        run_text = par.add_run(row['Questão'])
        run_text.font.size = Pt(12)

        alternativas = ['A', 'B', 'C', 'D']
        for alt in alternativas:
            doc.add_paragraph(f'({alt}) {row[alt]}')

        doc.add_paragraph('')  # Espaço entre questões

    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def montar_gabarito_doc(df, titulo):
    data = date.today().strftime("%d/%m/%y")

    doc = Document()
    doc.add_heading(titulo, 0)

    # Subtítulo com estilo 'Subtitle'
    sub = doc.add_paragraph(data)
    sub.style = 'Subtitle'

    lista = list(df)
    for row in range(len(lista)):
        # Cria o parágrafo
        par = doc.add_paragraph()

        # Adiciona o número em negrito
        run_num = par.add_run(f"{row + 1}. ")
        run_num.bold = True
        run_num.font.size = Pt(12)

        # Adiciona o texto da questão normal
        run_text = par.add_run(df.iloc[row])
        run_text.font.size = Pt(12)

        #doc.add_paragraph('')  # Espaço entre questões

    for p in doc.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer